#!/usr/bin/env python
# coding=utf-8
'''
作    者 : 北极星光 light22@126.com
创建时间 : 2024-06-02 00:21:43
最后修改 : 2024-11-10 01:36:18
修 改 者 : 北极星光
'''

import os
import datetime
from docx import Document
from openpyxl import load_workbook
from tools import *

# 生成结算申报资料


def generate_settlement_report():
    # 读取config.xlsx文件
    workbook = load_workbook('config.xlsx', data_only=True)
    sheet = workbook['结算申报']

    # 遍历sheet
    for i in range(2, sheet.max_row + 1):
        protocol_name = sheet[f'A{i}'].value  # 合同名称 <protocol_name>
        contract_number = sheet[f'B{i}'].value  # 合同编号 <contract_number>
        project_name = sheet[f'C{i}'].value  # 工程名称 <project_name>
        party_a_name = sheet[f'D{i}'].value  # 甲方名称 <party_a_name>
        party_b_name = sheet[f'E{i}'].value  # 乙方名称 <party_b_name>
        settlement_declaration_date = excel_date_format(sheet[f'F{i}'].value).strftime('%Y年%m月%d日')  # 结算申报日期 <settlement_declaration_date>
        settlement_declaration_year_and_month = excel_date_format(sheet[f'F{i}'].value).strftime('%Y年%m月')  # 结算申报年月 <settlement_declaration_year_and_month>
        contract_amount = round(sheet[f'G{i}'].value, 2)  # 合同金额 <contract_amount>
        settlement_amount = round(sheet[f'H{i}'].value, 2)  # 结算金额 <settlement_amount>
        settlement_amount_capital = num_to_capital(settlement_amount)  # 结算金额大写
        paid_amount = round(sheet[f'I{i}'].value, 2)  # 已支付金额 <paid_amount>
        paid_invoice_amount = round(sheet[f'J{i}'].value, 2)  # 已开发票金额 <paid_invoice_amount>
      
        # 生成 1结算--封面
        document = Document('data/1结算--封面.docx')
        # 替换模板中的变量
        replace_variables(
            document,
            project_name=project_name,  # 工程名称 <project_name>
            party_a_name=party_a_name,  # 甲方名称 <party_a_name>
            party_b_name=party_b_name,  # 乙方名称 <party_b_name>
            settlement_declaration_year_and_month=settlement_declaration_year_and_month  # 结算申报年月 <settlement_declaration_year_and_month>
        )

        # 创建工程文件夹
        os.makedirs(f'output/{project_name}工程结算申报资料', exist_ok=True)

        # 保存文件
        document.save(f'output/{project_name}工程结算申报资料/1结算--封面.docx')
        print(f'生成{project_name}工程 1结算--封面 成功！')

        # 生成 2结算--目录
        document = Document('data/2结算--目录.docx')
        # 替换模板中的变量
        replace_variables(
            document,
            project_name=project_name,  # 工程名称 <project_name>
        )
        # 保存文件
        document.save(f'output/{project_name}工程结算申报资料/2结算--目录.docx')
        print(f'生成{project_name}工程 2结算--目录 成功！')

        # 生成 3结算--工程结算申请
        document = Document('data/3结算--工程结算申请.docx')
        # 替换模板中的变量
        replace_variables(
            document,
            protocol_name=protocol_name,  # 合同名称 <protocol_name>
            contract_number=contract_number,  # 合同编号 <contract_number>
            party_a_name=party_a_name,  # 甲方名称 <party_a_name>
            settlement_amount=settlement_amount,  # 结算金额 <settlement_amount>
            settlement_amount_capital=settlement_amount_capital,  # 结算金额大写 <settlement_amount_capital>
            settlement_declaration_date=settlement_declaration_date,  # 结算申报日期 <settlement_declaration_date>
        )
        # 保存文件
        document.save(f'output/{project_name}工程结算申报资料/3结算--工程结算申请.docx')
        print(f'生成{project_name}工程 3结算--工程结算申请 成功！')

        # 生成 4结算--编制说明及结算汇总表
        wb = load_workbook('data/4结算--编制说明及结算汇总表.xlsx')
        # 替换模板中的变量
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            replace_variables_sheet(
                ws,
                protocol_name=protocol_name,  # 合同名称 <protocol_name>
                project_name=project_name,  # 工程名称 <project_name>
                contract_number=contract_number,  # 合同编号 <contract_number>
                contract_amount=contract_amount,  # 合同金额 <contract_amount>
                party_b_name=party_b_name,  # 乙方名称 <party_b_name>
            )
        # 保存文件
        wb.save(f'output/{project_name}工程结算申报资料/4结算--编制说明及结算汇总表.xlsx')
        print(f'生成{project_name}工程 4结算--编制说明及结算汇总表 成功！')

        # 生成 8结算--已付款统计清单
        wb = load_workbook('data/8结算--已付款统计清单.xlsx')
        # 替换模板中的变量
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            replace_variables_sheet(
                ws,
                protocol_name=protocol_name,  # 框架协议名称 <protocol_name>
                paid_amount=paid_amount,  # 已支付金额 <paid_amount>
                paid_invoice_amount=paid_invoice_amount,  # 已开发票金额 <paid_invoice_amount>
            )
        # 保存文件
        wb.save(f'output/{project_name}工程结算申报资料/8结算--已付款统计清单.xlsx')
        print(f'生成{project_name}工程 8结算--已付款统计清单 成功！')
